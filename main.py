import requests
import re
import os 
from copy import copy
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import shutil


def get_url_user_and_check_it():
    while True: 
            url_user = input('Введите вашу ссылку: ')     
            try:
                response = requests.get(url_user)
            except requests.exceptions.MissingSchema:
                print('ссылка неправильно введена')
                continue
            if "wikipedia.org" not in url_user:
                print('Введите ссылку на википедию')
                continue
            if response.status_code != 200:
                print('неизвестная ошибка сервера.Возможно неверно введена ссылка')
                continue 
            return response

def setting_styles(doc):
    style = doc.styles['Normal']
    style.font.size = Pt(14)
    style.font.name =  "Times New Roman"
    return style

def fill_in_file(all_tags,doc):  
    count = 0    
    for tag in all_tags: 
        real_tag = BeautifulSoup(tag,features="html.parser")
        if  'mw-file-description' in tag:
            count += 1
            img_link = real_tag.find('img').get('src')
            response = requests.get(f'https:{img_link}')
            response.raise_for_status()
            p = doc.add_paragraph()
            with open(f'img/picture{count}.jpg', "wb") as file:
                file.write(response.content)
            run = p.add_run()
            run.add_picture(image_to_jpg(f'img/picture{count}.jpg'))
            p.alignment =  WD_ALIGN_PARAGRAPH.CENTER
            delete_everything_in_folder('img')
        if '<p>' in tag:
            p = doc.add_paragraph(real_tag.text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue
        for a in range(1,5):
            if f'<h{a}' in tag:
                head = doc.add_heading(real_tag.text, level=a)
                head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
def get_rid_of_unnecessary_sections_bottom(all_tags):   
    reversed_all_tags = list(reversed(all_tags))
    for i in reversed_all_tags:
        if '<p>' in str(i):
            index = reversed_all_tags.index(i)
            del reversed_all_tags[0:index]
            return list(reversed(reversed_all_tags))

def get_tags_without_notes(all_tags):
        return [re.sub(r"\[.*?\]",r"", str(tag)) for tag in all_tags]

def try_save(doc,file_name):
        try:
            doc.save(f'{file_name}.docx')
        except PermissionError:
            print('Не удалось сохранить файл ')

def image_to_jpg(image_path):
    Image.open(image_path).convert('RGB').save(image_path)
    return image_path

def delete_everything_in_folder(folder_path):
                shutil.rmtree(folder_path)
                os.mkdir(folder_path)

if __name__ == '__main__':
    response = get_url_user_and_check_it()

    soup = BeautifulSoup(response.text, features="html.parser")
    mw_content_text_block = soup.find('div',id = 'mw-content-text')
    
    
    doc = Document()
    
    main_block = soup.find('div',class_="mw-content-ltr mw-parser-output")
    
    
    all_tags = main_block.find_all(['p', 'h1', 'h2', 'h3', 'h4','a'])
    h1 =  soup.find('h1', class_="firstHeading mw-first-heading").text
    
    styles = setting_styles(doc)
    head = doc.add_heading(h1,level =1)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    all_tags = get_rid_of_unnecessary_sections_bottom(all_tags)
    # all_tags = get_rid_of_unnecessary_sections_from_above(all_tags)
    all_tags = get_tags_without_notes(all_tags)
    fill_in_file(all_tags,doc)
    file_name = input('Выберите имя файла: ')
    try_save(doc,file_name)
    
